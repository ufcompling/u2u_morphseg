# Used Copilot to autofill and debug.
import io
import os
import json
from typing import Union

def save_text(file_path: str, data) -> None:
    """Saves a file to the /data directory in the Pyodide virtual filesystem.

    Accepts bytes, str, bytearray, memoryview, or lists of ints. Strings are
    encoded as UTF-8 before writing. This keeps the Python API simple for
    callers coming from JS.
    Ensures parent directory exists before writing.
    """
    # Ensure parent directory exists
    import os
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    # Normalize incoming data to bytes
    if isinstance(data, bytes):
        data_bytes = data
    elif isinstance(data, str):
        data_bytes = data.encode('utf-8')
    elif isinstance(data, bytearray):
        data_bytes = bytes(data)
    elif isinstance(data, memoryview):
        data_bytes = data.tobytes()
    elif isinstance(data, list):
        data_bytes = bytes(data)
    else:
        raise ValueError('Unsupported data type for saving file')
    with open(file_path, 'wb') as f:
        f.write(data_bytes)

def save_binary(file_name: str, data) -> None:
    """Save the bytes to string. Ensures parent directory exists before writing."""
    import os
    try:
        data_bytes = bytes(data)
    except Exception:
        raise ValueError('Data must be bytish for save_binary')
    file_path = file_name if file_name.startswith('/') else os.path.join('/data', file_name)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    with open(file_path, 'wb') as f:
        f.write(data_bytes)

def get_snapshot(dir: str) -> str:
    """Creates a serialized snapshot of the specified directory defualting to /data. Returns a JSON string with the file names and their contents as lists of bytes."""
    if dir is None:
        dir = '/data'
    snapshot = {}
    if os.path.exists(dir):
        for file_name in os.listdir(dir):
            if file_name == "crf.model":
                continue  # Skip the model file
            file_path = os.path.join(dir, file_name)
            if os.path.isfile(file_path):
                with open(file_path, 'rb') as f:
                    snapshot[file_name] = list(f.read())
    return json.dumps(snapshot)

def read_snapshot(snapshot_json: str, dir: str) -> None:
    """Reads a serialized snapshot JSON string and writes the files to the specified directory."""
    if dir is None:
        dir = '/data'
    snapshot = json.loads(snapshot_json)
    for file_name, content in snapshot.items():
        file_path = os.path.join(dir, file_name)
        save_binary(file_path, bytes(content))


def create_pdf(file_name: str, file_content: str) -> None:
    """Creates a PDF file with only the processed text as content, using default formatting."""
    from reportlab.pdfgen import canvas  # type: ignore
    from reportlab.lib.pagesizes import letter  # type: ignore
    import os
    new_name = file_name.rsplit('.', 1)[0] + '_processed.pdf' if '.' in file_name else file_name + '_processed.pdf'
    processed_file = new_name if new_name.startswith('/') else os.path.join('/data', new_name)
    file = canvas.Canvas(processed_file, pagesize=letter)
    file.setFont("Helvetica", 12)
    top_margin = 750
    left_margin = 40
    line_height = 16
    bottom_margin = 40
    y = top_margin
    lines = file_content.splitlines()
    if not lines:
        lines = ["(No content)"]
    for line in lines:
        file.drawString(left_margin, y, line)
        y -= line_height
        # Only create a new page if needed
        if y < bottom_margin and line != lines[-1]:
            file.showPage()
            file.setFont("Helvetica", 12)
            y = top_margin
    file.save()

def create_docx(file_name: str, file_content: str) -> None:
    """Creates a DOCX file hopefully with the original format."""
    from docx import Document # type: ignore
    old_file = file_name if file_name.startswith('/') else os.path.join('/data', file_name)
    new_name = file_name.rsplit('.', 1)[0] + '_processed.docx' if '.' in file_name else file_name + '_processed.docx'
    processed_file = new_name if new_name.startswith('/') else os.path.join('/data', new_name)
    # Load original DOCX
    doc = Document(old_file)
    # Remove all existing paragraphs
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
    for line in file_content.splitlines():
        doc.add_paragraph(line)
    doc.save(processed_file)

def read_file(file_path: str, detect_text: bool = True, encoding: str = 'utf-8') -> str:
    """Reads a file and returns a JSON string describing the result.

    Returns a JSON-encoded dict with two keys:
      - type: 'text' or 'Uint8Array'
      - content: decoded text (if 'text') or list of bytes (if 'Uint8Array')

    This avoids passing raw bytes across the JS/Py boundary and lets the JS
    side decide how to display or download the data.
    """

    # Debug: print contents of /data and /data/English
    if not os.path.exists(file_path):
        # project.json/cycles.json/annotations.json are initialized lazily;
        lazy_files = ['/project.json', '/cycles.json', '/annotations.json']        
        if any(file_path.endswith(lazy) for lazy in lazy_files):
            return json.dumps({'type': 'text', 'content': ''})

        print(f"[db_worker] File not found: {file_path}")
        raise FileNotFoundError(f"File '{file_path}' not found in /data directory")

    with open(file_path, 'rb') as f:
        b = f.read()

    if detect_text:
        try:
            text = b.decode(encoding)
            return json.dumps({'type': 'text', 'content': text})
        except UnicodeDecodeError:
            pass

    # binary data as list of ints for JS to convert to Uint8Array
    return json.dumps({'type': 'Uint8Array', 'content': list(b)})


def delete_file(file_path: str) -> None:
    """Deletes a file from the /data directory in the Pyodide virtual filesystem."""
    if os.path.exists(file_path):
        os.remove(file_path)

def clear_files(directory: str | None = None) -> None:
    """Deletes all files in the specified directory. Defaults to /data."""
    if directory is None:
        directory = '/data'
    if os.path.exists(directory):
        for file_name in os.listdir(directory):
            file_path = os.path.join(directory, file_name)
            if os.path.isfile(file_path):
                os.remove(file_path)